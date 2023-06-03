# imagine you have unstructered data.
# like resumes
# let gpt search for some values, save them all in a dict and easy to clal them
# https://www.youtube.com/watch?v=2OqbpYpoq7A&ab_channel=LiamOttley
import ast
import json
import re
from constants import OPENAI_API_KEY

import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
import openai

openai.api_key = OPENAI_API_KEY

def read_file(file):
    if file.name.endswith('.docx'):
        return read_docx(file)
    elif file.name.endswith('.pdf'):
        return read_pdf(file)
    else:
        print("Unsupported file format.")
        return None


def read_docx(file):
    doc = Document(file)
    text = ""

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"

    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            text += paragraph.text + "\n"

    return text

def get_dict(answer):
    try:
        dict_string = re.search(r'\{.*\}', answer).group()

        # Now 'dict_string' should be a string that looks like a dictionary.
        # You can use the json.loads function to convert it to an actual dictionary:
        dict_from_user = json.loads(dict_string)
    except:
        # Find string representation of dictionary using regular expression
        dict_from_user = re.search('```(.*?)```', answer, re.DOTALL).group(1).strip()
        # Replace single quotes with double quotes to prepare string for json.loads
        dict_from_user = dict_from_user.replace("'", '"')

    return dict_from_user

def get_data_keys(dict, doc, model="gpt-3.5-turbo"):
    messages.append({"role": "system", "content": "You are a helpful assistant who is extreme good at searching for information in a file and returning it very need and constant in a dict that is provided to you."})
    messages.append({"role": "user", "content": f"Given the following dict: {dict}, please analyze the provided resume and populate the keys with the relevant information found. If no information is found for a key, leave the value as None. If multiple pieces of information correspond to a single key, please collate these into a list. Only look for the keys in the dict and dont add anything. Upon completion, please return the dict with the same keys and the values you added. {doc}"})
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.1,
    )
    answer = response.choices[0].message["content"]
    return answer

def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

data_keys = {'name': None}
update_data_keys = []
text_input = {}
num = 1
messages = []

# Create a title for the application
st.title('File Upload and Text Input App')

# Create file uploader
uploaded_files = st.file_uploader("Choose a file", accept_multiple_files=True)

# Iterate over the uploaded files and display file details
for uploaded_file in uploaded_files:
    text = read_file(uploaded_file)
    if text == "":
        st.write(f"could not read {uploaded_file}")
        continue
    text_input.update({num: text})
    num += 1


if "textbox_count" not in st.session_state:
    st.session_state.textbox_count = 1

if st.button('Add data key'):
    st.session_state.textbox_count += 1

for i in range(st.session_state.textbox_count):
    text = st.text_input(f'Data key {i+1}')
    data_keys.update({text: None})

if st.button('Submit'):
    print(text_input)
    for resume in text_input.values():
        answer = get_data_keys(data_keys, resume)
        print(answer)
        dict_from_user = get_dict(answer)
        update_data_keys.append(dict_from_user)
    st.write(update_data_keys)


